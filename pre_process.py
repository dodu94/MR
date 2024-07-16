import re
import calendar
from datetime import datetime, timedelta
import docx
import io
from docx.table import _Cell
from copy import deepcopy
from process_dms import parse_DMS
import logging

# PARAMETERS
REPORT_NUMBER = {"table": 0, "cell": (0, 2)}
DMS_CELL = {"table": 0, "cell": (2, 2)}
H_IN_PERIOD_CELL = {"table": 1, "cell": (1, 3)}
PERIODS = [
    {"table": 2, "cell": (0, 1)},
    {"table": 3, "cell": (0, 1)},
    {"table": 4, "cell": (0, 1)},
]
AUTHOR_NAME = {"table": 6, "cell": (1, 0)}
DATE_AUTHOR = {"table": 6, "cell": (2, 0)}
DATE_APPROVAL = {"table": 6, "cell": (2, 1)}
NEW_MILESTONE = {"table": 2, "cell": (1, 1)}
MILESTONE_TO_COPY = {"table": 4, "cell": (1, 1)}
TO_HIGHLIGHT = [
    {"table": 1, "cell": (1, 3)},
    {"table": 3, "cell": (1, 1)},
    {"table": 5, "cell": (1, 1)},
    MILESTONE_TO_COPY,
]

MONTH_NAMES = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}

# PATTERNS
month_number_pat = re.compile(r"M\d+")
date_pat = re.compile(r"(\d{2})/(\d{2})/(\d{4})")
period_pat = re.compile(r"\d{2}/\d{2}/\d{4}\s*[-–]\s*\d{2}/\d{2}/\d{4}")
pat_file_name = re.compile(r"#\d+\s+M\d+\s+\d+")


# FUNCTIONS
def make_cell_bold(cell: _Cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


def copy_text_body(cell_src: _Cell, cell_dest: _Cell) -> None:
    new_paragraphs = []
    for p in cell_src.paragraphs:
        inserted_p = deepcopy(p._p)
        if p._p.get_or_add_pPr().numPr:
            inserted_p.style = "ListNumber"
        new_paragraphs.append(inserted_p)
    cell_dest._element.clear_content()
    for p in new_paragraphs:
        cell_dest._element.append(p)


def change_period_str(period_str: str) -> str:
    try:
        number = month_number_pat.search(period_str).group()
    except AttributeError:
        raise RuntimeError(f"{period_str} does not contain a month number (e.g. M02).")
    int_number = int(number[1:])
    # handle the case where the month number is 12
    if int_number == 12:
        new_str = period_str.replace(number, "M01")
    else:
        new_str = period_str.replace(number, f"M{int_number + 1:02d}")
    return new_str


def highlight_cell(cell: _Cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW


def change_period_date(period_date: str) -> str:
    try:
        day, month, year = date_pat.search(period_date).group().split("/")
    except AttributeError:
        return period_date
    # get day interval for next month
    # if last month was december, the new year needs to be increased by 1
    if int(month) == 12:
        year = int(year) + 1
        month = 1
    else:
        month = int(month) + 1
        year = int(year)
    last_day = calendar.monthrange(year, month)[1]

    new_period = f"01/{month:02d}/{year} - {last_day:02d}/{month:02d}/{year}"
    return period_pat.sub(new_period, period_date)


def update_report_number(num_str: str) -> str:
    num, month, year = num_str.split("_")
    return _update_report_number(num, month, year)


def _update_report_number(num: str, month: str, year: str):
    newmonth = change_period_str(month)
    if newmonth == "M01":
        year = int(year) + 1
    new_string_name = f"#{int(num[1:])+1:02d}_{newmonth}_{year}"
    newmonth = MONTH_NAMES[int(newmonth[1:])]

    return new_string_name, newmonth, year


def pre_process(file: io.BytesIO) -> tuple[docx.Document, str]:
    # Identify the old monthly report
    # file = None
    # for file in os.listdir():
    #     if file.endswith(".docx"):
    #         old_mr = docx.Document(file)
    #         break
    # if file is None:
    #     raise FileNotFoundError("No .docx file found in the current directory")
    old_mr = docx.Document(file)

    # -- Update the month number and period --
    for period in PERIODS:
        cell = old_mr.tables[period["table"]].cell(*period["cell"])
        old_text = cell.text
        new_text = change_period_date(old_text)
        new_text = change_period_str(new_text)
        cell.text = new_text
        make_cell_bold(cell)

    # -- Update the signing date --
    # Get today's date and tomorrow's date
    today = datetime.now().date()
    # tomorrow = today + timedelta(days=1)

    date_author = old_mr.tables[DATE_AUTHOR["table"]].cell(*DATE_AUTHOR["cell"])
    date_author.text = "Date: {:%d/%m/%Y}".format(today)
    date_approval = old_mr.tables[DATE_APPROVAL["table"]].cell(*DATE_APPROVAL["cell"])
    date_approval.text = "Date: {:%d/%m/%Y}".format(today)

    # -- Update the report number --
    report_number = old_mr.tables[REPORT_NUMBER["table"]].cell(*REPORT_NUMBER["cell"])
    report_number.text = update_report_number(report_number.text)

    # -- Update the new milestone section --
    new_milestone = old_mr.tables[NEW_MILESTONE["table"]].cell(*NEW_MILESTONE["cell"])
    milestone_to_copy = old_mr.tables[MILESTONE_TO_COPY["table"]].cell(
        *MILESTONE_TO_COPY["cell"]
    )
    copy_text_body(milestone_to_copy, new_milestone)

    # -- Highlight the cells that need to be still modified --
    for location in TO_HIGHLIGHT:
        cell = old_mr.tables[location["table"]].cell(*location["cell"])
        highlight_cell(cell)

    # -- get the new file name --
    match = pat_file_name.search(file.name)
    num, month, year = match.group().split()
    new_number, newmonth, newyear = _update_report_number(num, month, year).replace(
        "_", " "
    )
    new_name = file.name[: match.start()] + new_number + ".docx"

    # --- try to get name and dms ---
    dms_numbers = parse_DMS()
    author_name = (
        old_mr.tables[AUTHOR_NAME["table"]]
        .cell(*AUTHOR_NAME["cell"])
        .text.split("/")
        .strip()
    )
    try:
        dms = dms_numbers.loc[author_name, newyear, newmonth]
    except:
        logging.error(
            f"Could not find DMS number for {author_name} in {newyear} {newmonth}"
        )
        dms = "CHANGE THIS, DMS NOT FOUND"
    old_mr.tables[DMS_CELL["table"]].cell(*DMS_CELL["cell"]).text = dms

    # save to stream
    new_mr = io.BytesIO()
    old_mr.save(new_mr)

    return new_mr, new_name
